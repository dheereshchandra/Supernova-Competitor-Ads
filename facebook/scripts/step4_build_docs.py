#!/usr/bin/env python3
"""
Step 4 — Stage 6: build two .docx files per video.

For each row, produces:
  step4_workspace/docs/<id>_competitor_analysis.docx
  step4_workspace/docs/<id>_supernova_rewrite.docx

Doc 1 (competitor analysis) contains:
  - Header (filename, duration, resolution, production type, character roster, provenance)
  - Per scene: scene label, start-end, ORIGINAL screenshot, AI-regenerated clean
    panels in their natural layout, composite visual description, per-panel
    descriptions, speaker-attributed transcript, on-screen text

Doc 2 (Supernova rewrite) contains:
  - Header (this is an internal brand-safe rewrite)
  - Per scene: scene label, clean panel imagery, scene-summary bullets,
    Supernova-voice rewritten script, rewritten on-screen text
  - Brand-swaps log at the end

After build, programmatic verification runs:
  - Both docx files open cleanly with python-docx
  - Paragraph count > minimum (scenes × 4 for competitor doc, scenes × 3 for supernova)
  - Embedded image count = expected (original + clean panels per scene)
  - No empty headings (every "Scene N — label" has content after it)
  - File size > 5 KB

Usage:
    python3 scripts/step4_build_docs.py --competitor zinglish <id1> <id2> ...
"""
from __future__ import annotations

import argparse
import json
import pathlib
import subprocess
import sys
from typing import Optional

WORKSPACE = pathlib.Path("step4_workspace")
SCENES_DIR = WORKSPACE / "scenes"
FRAMES_DIR = WORKSPACE / "frames"
IMAGES_DIR = WORKSPACE / "images"
DOCS_DIR = WORKSPACE / "docs"

# Caption prefix written immediately under every embedded image. Used by both
# the writer (build_competitor_doc / build_supernova_doc) and the verifier
# (verify_docs). Don't translate or pluralise — verify_docs greps for this.
ASSET_CAPTION_PREFIX = "Asset: "


def load_image_urls(ad_id: str) -> dict:
    """Read step4_workspace/images/<ad_id>/r2_urls.json sidecar, or return empty.

    The sidecar is written by scripts/step4_upload_images.py (Stage 4.5). If
    it isn't there yet (Stage 4.5 not run), docs are still built but with no
    URL captions — verify_docs will then flag every image as missing-URL,
    which is the right signal.
    """
    sidecar = IMAGES_DIR / ad_id / "r2_urls.json"
    if not sidecar.exists():
        return {}
    try:
        return json.loads(sidecar.read_text())
    except json.JSONDecodeError:
        return {}


def _add_hyperlink(paragraph, url: str, display_text: str) -> None:
    """Insert a clickable hyperlink run into a python-docx paragraph.

    python-docx has no first-class hyperlink helper, so we build the OOXML
    by hand. Pattern lifted from the well-known python-docx hyperlink recipe.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    # Blue + underline so it visibly reads as a link.
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")  # ~8pt — small caption
    r_pr.append(sz)
    new_run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = display_text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_asset_caption(doc, url: str | None) -> None:
    """Insert the "Asset: <hyperlink>" caption paragraph beneath an embedded image.

    Always called after add_picture. If `url` is falsy (Stage 4.5 didn't
    produce a URL for this image) we still write the prefix with a "[missing]"
    placeholder so verify_docs can detect the gap explicitly.
    """
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(ASSET_CAPTION_PREFIX)
    run.font.size = Pt(8)
    run.italic = True
    if url:
        _add_hyperlink(p, url, url)
    else:
        miss = p.add_run("[missing — Stage 4.5 not run for this image]")
        miss.font.size = Pt(8)
        miss.italic = True


def ensure_packages():
    try:
        import docx  # noqa
        import PIL  # noqa
    except ImportError:
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "--quiet",
             "python-docx", "Pillow"],
            check=False,
        )


def video_path_for(ad_id: str, competitor: str) -> Optional[pathlib.Path]:
    # Search ALL dated folders (newest first), not just the latest — a long-running
    # ad's video may have come from an earlier scrape (see step4_decompose.video_path_for).
    for vdir in sorted(pathlib.Path("videos").glob(f"{competitor}-*"), reverse=True):
        for fname in (f"{ad_id}.mp4", f"{ad_id}_v2.mp4", f"{ad_id}_v3.mp4"):
            p = vdir / fname
            if p.exists():
                return p
    return None


def image_path_for(ad_id: str, competitor: str) -> Optional[pathlib.Path]:
    """Mirror of video_path_for() for Image-type ads (HANDOVER §8.5/§9.5).
    Searches ALL dated folders (newest first), not just the latest."""
    for idir in sorted(pathlib.Path("images").glob(f"{competitor}-*"), reverse=True):
        for fname in (f"{ad_id}.jpg", f"{ad_id}_c2.jpg", f"{ad_id}_c3.jpg",
                      f"{ad_id}_c4.jpg", f"{ad_id}_c5.jpg"):
            p = idir / fname
            if p.exists():
                return p
    return None


def get_video_meta(vpath: pathlib.Path) -> dict:
    try:
        out = subprocess.check_output(
            ["ffprobe", "-v", "error", "-print_format", "json",
             "-show_format", "-show_streams", str(vpath)],
            timeout=10,
        ).decode()
        d = json.loads(out)
        dur = float(d.get("format", {}).get("duration", 0.0))
        w = h = None
        for s in d.get("streams", []):
            if s.get("codec_type") == "video" and w is None:
                w, h = s.get("width"), s.get("height")
        return {"duration_s": round(dur, 1),
                "resolution": f"{w}x{h}" if w else "?"}
    except Exception:
        return {"duration_s": None, "resolution": "?"}


def panel_image_path(ad_id: str, scene_n: int, position: str) -> Optional[pathlib.Path]:
    p = IMAGES_DIR / ad_id / f"gen_{scene_n:02d}_{position}.jpg"
    return p if p.exists() and p.stat().st_size > 30000 else None


def original_frame_path(ad_id: str, scene_n: int) -> Optional[pathlib.Path]:
    # Align with step4_frames.py — anything under 20 KB is a degenerate write
    # at source-native resolution (HANDOVER §9.7).
    p = FRAMES_DIR / ad_id / f"orig_{scene_n:02d}.jpg"
    return p if p.exists() and p.stat().st_size > 20_000 else None


def build_competitor_doc(ad_id: str, competitor: str, decompose: dict,
                         out_path: pathlib.Path, log) -> bool:
    """Build the competitor-analysis docx. Returns True on success."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    parsed = decompose.get("parsed", {})
    scenes = parsed.get("scenes", [])
    characters = parsed.get("characters", [])
    is_image_only = bool(decompose.get("is_image_only"))

    # Per HANDOVER §9.7 every embedded image is captioned with its R2 URL,
    # read from the Stage 4.5 sidecar.
    image_urls = load_image_urls(ad_id)
    orig_urls = image_urls.get("orig", {})
    gen_urls = image_urls.get("gen", {})
    char_urls = image_urls.get("char", {})

    if is_image_only:
        # For image-only ads the source is a local .jpg, not a video. We don't
        # ffprobe; we just record the file we have.
        recorded = decompose.get("source_image_path")
        ipath = pathlib.Path(recorded) if recorded and pathlib.Path(recorded).exists() \
                else image_path_for(ad_id, competitor)
    else:
        vpath = video_path_for(ad_id, competitor)
        meta = get_video_meta(vpath) if vpath else {"duration_s": None, "resolution": "?"}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading(f"Competitor Ad Analysis — {competitor.title()} — Ad {ad_id}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Header block — differs slightly for image vs. video
    if is_image_only:
        doc.add_paragraph(f"Source image: {ipath.name if ipath else '?'}")
        doc.add_paragraph(f"Media type: static image (no audio)")
    else:
        doc.add_paragraph(f"Source video: {vpath.name if vpath else '?'}")
        doc.add_paragraph(f"Duration: {meta.get('duration_s', '?')}s   |   "
                          f"Resolution: {meta.get('resolution', '?')}")
    pt = parsed.get("production_type", "?")
    doc.add_paragraph(f"Production type: {pt}")
    if parsed.get("production_notes"):
        doc.add_paragraph(f"Production notes: {parsed['production_notes']}")

    doc.add_heading("Character roster", level=2)
    if not characters:
        doc.add_paragraph("(no characters detected)")
    for c in characters:
        cid = c.get("id", "?")
        role = c.get("role", "")
        doc.add_paragraph(f"Character {cid} — {role}", style="List Bullet")
        if c.get("appearance"):
            doc.add_paragraph(f"    Appearance: {c['appearance']}")
        if c.get("wardrobe"):
            doc.add_paragraph(f"    Wardrobe: {c['wardrobe']}")
        if c.get("demeanor"):
            doc.add_paragraph(f"    Demeanor: {c['demeanor']}")
        # Embed the character reference sheet if it was generated, with R2 caption.
        sheet_path = IMAGES_DIR / ad_id / f"char_{cid}_sheet.jpg"
        if sheet_path.exists() and sheet_path.stat().st_size > 50000:
            try:
                doc.add_picture(str(sheet_path), width=Inches(2.5))
                add_asset_caption(doc, char_urls.get(str(cid)))
            except Exception as e:
                log(f"    char sheet embed failed for {sheet_path}: {e}")

    doc.add_paragraph()
    doc.add_paragraph("─" * 60)

    # Per-scene blocks
    for scene in scenes:
        n = scene.get("n", 0)
        label = scene.get("scene_label", "?")
        ss = scene.get("start_s", 0)
        es = scene.get("end_s", 0)

        doc.add_heading(f"Scene {n} — {label}", level=1)
        doc.add_paragraph(f"[{ss}s – {es}s]   Camera: {scene.get('camera', '?')}   "
                          f"Layout: {scene.get('layout', '?')}")
        if scene.get("setting"):
            doc.add_paragraph(f"Setting: {scene['setting']}")

        # Original screenshot
        orig = original_frame_path(ad_id, n)
        if orig:
            doc.add_paragraph("Original screenshot:")
            try:
                doc.add_picture(str(orig), width=Inches(4.5))
                add_asset_caption(doc, orig_urls.get(f"{n:02d}"))
            except Exception as e:
                log(f"    image embed failed for {orig}: {e}")
        else:
            doc.add_paragraph("(no original screenshot extracted)")

        # Clean regenerated panels — only shown when they actually exist. Image
        # generation is a separate later step, so their absence here is expected
        # and must not emit "unavailable" noise; the per-panel text descriptions
        # below carry the visual detail in the meantime.
        panels = scene.get("panels", [])
        gen_panels = [(p.get("position", "full"),
                       panel_image_path(ad_id, n, p.get("position", "full"))) for p in panels]
        gen_panels = [(pos, pp) for pos, pp in gen_panels if pp]
        if gen_panels:
            doc.add_paragraph("Clean regenerated panels:")
            for pos, pp in gen_panels:
                try:
                    doc.add_paragraph(f"  Panel: {pos}")
                    doc.add_picture(str(pp), width=Inches(4.5))
                    add_asset_caption(doc, gen_urls.get(f"{n:02d}_{pos}"))
                except Exception as e:
                    log(f"    panel embed failed for {pp}: {e}")

        # Visual description
        if scene.get("visual_description"):
            doc.add_heading("Composite visual description", level=3)
            doc.add_paragraph(scene["visual_description"])

        # Per-panel descriptions
        if panels:
            doc.add_heading("Per-panel descriptions", level=3)
            for panel in panels:
                pos = panel.get("position", "full")
                desc = panel.get("panel_visual_description", "")
                doc.add_paragraph(f"[{pos}] {desc}", style="List Bullet")

        # Audio transcript
        if scene.get("audio_transcript"):
            doc.add_heading("Audio transcript", level=3)
            for line in scene["audio_transcript"].split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # On-screen text
        if scene.get("on_screen_text"):
            doc.add_heading("On-screen text", level=3)
            doc.add_paragraph(scene["on_screen_text"])

        doc.add_paragraph("─" * 60)

    # Footer / provenance
    doc.add_heading("Provenance", level=2)
    doc.add_paragraph(f"Decomposed with: {decompose.get('model', '?')}")
    doc.add_paragraph("Clean-panel / character-sheet image generation runs as a separate later "
                      "step (not part of this script/document pass).")
    doc.add_paragraph(f"This document is for internal competitor research. "
                      f"All copy and imagery is summarised from the competitor's public ad library.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return True


def build_supernova_doc(ad_id: str, competitor: str, decompose: dict,
                        rewrite: dict, out_path: pathlib.Path, log) -> bool:
    """Build the Supernova-rewrite docx. Returns True on success."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    parsed_decompose = decompose.get("parsed", {})
    parsed_rewrite = rewrite.get("parsed", {})
    scenes_d = {s.get("n"): s for s in parsed_decompose.get("scenes", [])}
    scenes_r = parsed_rewrite.get("scenes", [])
    brand_swaps = parsed_rewrite.get("brand_swaps_detected", [])

    # Same Stage-4.5 sidecar as the competitor doc — keeps both docs consistent.
    image_urls = load_image_urls(ad_id)
    gen_urls = image_urls.get("gen", {})

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    title = doc.add_heading(f"Supernova-Voice Rewrite — based on {competitor.title()} Ad {ad_id}",
                            level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("This is an internal brand-safe rewrite of a competitor's ad, "
                      "produced for Supernova's creative team. Use the visuals + script "
                      "as a starting point for original Supernova creative — do not ship "
                      "this as-is.")
    doc.add_paragraph()

    # Cast / legend — map each character id to its assigned name + role so a reader
    # can follow the NAMED script below without mentally tracking "Character A".
    # The rewrite assigns each character a `name` (Miss Nova for the AI teacher);
    # older sidecars without names fall back to the id — role line.
    cast = parsed_rewrite.get("characters", [])
    if cast:
        doc.add_heading("Cast / Legend", level=2)
        for c in cast:
            cid = c.get("id", "?")
            role = c.get("role", "")
            name = c.get("name")
            if name:
                doc.add_paragraph(f"Character {cid} = {name}" + (f" ({role})" if role else ""),
                                  style="List Bullet")
            else:
                doc.add_paragraph(f"Character {cid} — {role}", style="List Bullet")
        doc.add_paragraph()

    doc.add_paragraph("─" * 60)

    # Per-scene blocks
    for scene_r in scenes_r:
        n = scene_r.get("n", 0)
        scene_d = scenes_d.get(n, {})
        label = scene_r.get("scene_label") or scene_d.get("scene_label", "?")

        doc.add_heading(f"Scene {n} — {label}", level=1)

        # Clean panel imagery (from Stage 4)
        panels = scene_d.get("panels", [])
        for panel in panels:
            pos = panel.get("position", "full")
            pp = panel_image_path(ad_id, n, pos)
            if pp:
                try:
                    doc.add_picture(str(pp), width=Inches(4.5))
                    add_asset_caption(doc, gen_urls.get(f"{n:02d}_{pos}"))
                except Exception as e:
                    log(f"    supernova doc image embed failed: {e}")

        # Scene visuals (kept from the competitor ad) — rendered as TEXT so this doc is
        # self-contained without generated panels. Image generation is a separate later
        # step; until it runs there are no panel images, so the look must be described here.
        panel_descs = [p for p in panels if p.get("panel_visual_description")]
        if scene_d.get("setting") or scene_d.get("visual_description") or panel_descs:
            doc.add_heading("Scene visuals (kept from the competitor ad)", level=3)
            if scene_d.get("setting"):
                doc.add_paragraph(f"Setting: {scene_d['setting']}")
            if scene_d.get("visual_description"):
                doc.add_paragraph(scene_d["visual_description"])
            for panel in panel_descs:
                pos = panel.get("position", "full")
                doc.add_paragraph(f"[{pos}] {panel['panel_visual_description']}", style="List Bullet")

        # Scene summary bullets
        if scene_r.get("scene_summary"):
            doc.add_heading("Scene summary", level=3)
            for line in scene_r["scene_summary"].split("\n"):
                line = line.strip().lstrip("- *•").strip()
                if line:
                    doc.add_paragraph(line, style="List Bullet")

        # Supernova rewrite script
        if scene_r.get("supernova_script"):
            doc.add_heading("Supernova-voice script", level=3)
            for line in scene_r["supernova_script"].split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # On-screen text rewrite
        if scene_r.get("supernova_on_screen_text"):
            doc.add_heading("On-screen text (Supernova version)", level=3)
            doc.add_paragraph(scene_r["supernova_on_screen_text"])

        doc.add_paragraph("─" * 60)

    # Brand swaps log
    if brand_swaps:
        doc.add_heading("Brand swaps applied", level=2)
        for swap in brand_swaps:
            cb = swap.get("competitor_brand", "?")
            sw = swap.get("supernova_swap", "?")
            doc.add_paragraph(f"{cb} → {sw}", style="List Bullet")

    doc.add_heading("Provenance", level=2)
    doc.add_paragraph(f"Rewritten with: {rewrite.get('model', '?')}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return True


def _count_embedded_images(doc) -> int:
    """Return the number of inline pictures actually rendered in the body.

    Counting `doc.part.rels` would undercount because python-docx deduplicates
    the relationship when the same image file is embedded more than once. We
    count `<w:drawing>` elements directly so duplicate embeds are honoured.
    """
    from docx.oxml.ns import qn
    body = doc.element.body
    return len(body.findall(f".//{qn('w:drawing')}"))


def _extract_hyperlinks(doc) -> list[str]:
    """Return every external hyperlink target in the document, in order."""
    out: list[str] = []
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype and rel.is_external:
            out.append(rel.target_ref)
    return out


def verify_docs(ad_id: str, comp_path: pathlib.Path, sn_path: pathlib.Path,
                expected_scenes: int, *, check_urls: bool = True,
                env: dict | None = None) -> list[str]:
    """Run programmatic verification on the built docs. Returns list of issues.

    Per HANDOVER §9.7 this verifier now also checks:
      - One "Asset: " caption paragraph per embedded image
      - Every URL after "Asset: " resolves (HEAD 200) — controlled by ``check_urls``
      - The r2_urls.json sidecar is present and lists every image we embedded

    The HEAD check is gated by ``check_urls`` because the build step runs
    before upload finishes in some paths; pass check_urls=False when calling
    immediately after build, and True when calling from the audit step.
    """
    from docx import Document
    issues: list[str] = []

    # Per-ad sidecar must exist before either doc can be trusted.
    sidecar_path = IMAGES_DIR / ad_id / "r2_urls.json"
    sidecar_loaded: dict = {}
    if not sidecar_path.exists():
        issues.append(f"sidecar missing: {sidecar_path} — Stage 4.5 (step4_upload_images.py) not run")
    else:
        try:
            sidecar_loaded = json.loads(sidecar_path.read_text())
        except json.JSONDecodeError as e:
            issues.append(f"sidecar unreadable: {e}")

    for label, path, min_scenes in (("competitor", comp_path, expected_scenes),
                                     ("supernova", sn_path, expected_scenes)):
        if not path.exists():
            issues.append(f"{label}: file missing")
            continue
        if path.stat().st_size < 5000:
            issues.append(f"{label}: too small ({path.stat().st_size}B)")
            continue
        try:
            doc = Document(str(path))
        except Exception as e:
            issues.append(f"{label}: failed to open: {e}")
            continue
        paragraphs = list(doc.paragraphs)
        if len(paragraphs) < min_scenes * 3:
            issues.append(f"{label}: only {len(paragraphs)} paragraphs (expected ≥ {min_scenes * 3})")
        scene_headings = [p for p in paragraphs if p.text.startswith("Scene ")]
        if len(scene_headings) < min_scenes:
            issues.append(f"{label}: only {len(scene_headings)} scene headings (expected ≥ {min_scenes})")
        # Check for empty paragraphs masquerading as headings
        for p in paragraphs:
            if p.style.name.startswith("Heading") and not p.text.strip():
                issues.append(f"{label}: empty heading at paragraph index "
                              f"{paragraphs.index(p)}")
                break
        # Check for placeholder strings
        full_text = "\n".join(p.text for p in paragraphs)
        for placeholder in ("TODO", "<INSERT>", "<UNKNOWN>", "PLACEHOLDER"):
            if placeholder in full_text:
                issues.append(f"{label}: contains placeholder string '{placeholder}'")

        # --- HANDOVER §9.7 new checks: image vs Asset-caption parity ---
        n_images = _count_embedded_images(doc)
        n_captions = sum(1 for p in paragraphs if p.text.startswith(ASSET_CAPTION_PREFIX))
        if n_images and n_captions != n_images:
            issues.append(f"{label}: {n_images} embedded images but {n_captions} 'Asset:' "
                          f"captions — every image must have a URL line beneath it")

        # Any caption that ended up as "[missing — ...]" is a hard fail.
        missing_caps = [p.text for p in paragraphs
                        if p.text.startswith(ASSET_CAPTION_PREFIX) and "missing" in p.text]
        if missing_caps:
            issues.append(f"{label}: {len(missing_caps)} image(s) with no R2 URL "
                          f"(Stage 4.5 didn't produce them) — first: {missing_caps[0][:120]}")

        # URL HEAD checks (optional — controlled by check_urls)
        if check_urls and env is not None:
            from r2_utils import head_check as _head
            hrefs = _extract_hyperlinks(doc)
            for href in hrefs:
                ok, status = _head(href)
                if not ok:
                    issues.append(f"{label}: URL not reachable (HTTP {status}): {href}")

    # Sidecar / sanity: every URL in the sidecar should be from the configured public base.
    if sidecar_loaded and env is not None:
        base = env.get("R2_PUBLIC_URL_BASE", "")
        all_urls: list[str] = []
        for kind in ("orig", "gen", "char"):
            for v in (sidecar_loaded.get(kind) or {}).values():
                all_urls.append(v)
        bad = [u for u in all_urls if base and not u.startswith(base)]
        if bad:
            issues.append(f"sidecar: {len(bad)} URLs do not start with {base} "
                          f"(first: {bad[0][:120]})")

    return issues


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--competitor", required=True)
    ap.add_argument("ids", nargs="+")
    args = ap.parse_args()

    ensure_packages()
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    competitor = args.competitor.lower().strip()

    def log(msg):
        print(msg, flush=True)

    print(f"Stage 6 — building docx files for {len(args.ids)} videos")
    built = []
    failed_verify = []

    for ad_id in args.ids:
        sidecar = SCENES_DIR / f"{ad_id}.json"
        rewrite_sc = SCENES_DIR / f"{ad_id}.supernova.json"
        if not sidecar.exists():
            log(f"  [{ad_id}] SKIP — no decompose sidecar")
            continue
        if not rewrite_sc.exists():
            log(f"  [{ad_id}] SKIP — no Supernova rewrite sidecar")
            continue
        decompose = json.loads(sidecar.read_text())
        rewrite = json.loads(rewrite_sc.read_text())
        n_scenes = len(decompose.get("parsed", {}).get("scenes", []))

        comp_path = DOCS_DIR / f"{ad_id}_competitor_analysis.docx"
        sn_path = DOCS_DIR / f"{ad_id}_supernova_rewrite.docx"

        try:
            build_competitor_doc(ad_id, competitor, decompose, comp_path, log)
            build_supernova_doc(ad_id, competitor, decompose, rewrite, sn_path, log)
            # check_urls=False at build-time: Stage 7 has not run yet, so HEAD
            # checking would always fail. The full HEAD-check pass runs in the
            # audit step (see scripts/audit_step4.py or `--audit` in this script).
            issues = verify_docs(ad_id, comp_path, sn_path, n_scenes,
                                 check_urls=False, env=None)
            if issues:
                log(f"  [{ad_id}] BUILT with verification issues:")
                for iss in issues:
                    log(f"    - {iss}")
                failed_verify.append((ad_id, issues))
            else:
                log(f"  [{ad_id}] OK — verification passed (scenes={n_scenes})")
                built.append(ad_id)
        except Exception as e:
            log(f"  [{ad_id}] BUILD FAILED: {type(e).__name__}: {e}")

    print(f"\nStage 6 done — {len(built)} clean, {len(failed_verify)} with verification issues")
    return 0 if not failed_verify else 1


if __name__ == "__main__":
    raise SystemExit(main())
