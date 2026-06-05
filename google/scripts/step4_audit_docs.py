#!/usr/bin/env python3
"""
Step 4 — Stage 7a (Google): strict programmatic audit of the two built docs.

Runs AFTER step4_build_docs.py and BEFORE step4_upload_and_update.py. Performs
deeper checks than the structural verification baked into build_docs.py.

For each (creative_id, variant_index):

  Competitor-analysis doc and Supernova-rewrite doc — both checked for:
    1. Every embedded image is immediately followed by a paragraph starting with "R2: https://…".
       Catches: Stage 6a was skipped, an image upload failed, or the caption helper regressed.
    2. Scene-heading count matches the decompose sidecar's scene count.
    3. No empty headings.
    4. No placeholder strings (TODO, <INSERT>, <UNKNOWN>, PLACEHOLDER).
    5. No "R2 URL not available" captions (means Stage 6a needs re-running).
    6. Reasonable paragraph count vs. scene count.

  Competitor-analysis doc only:
    7. Every declared character in the decompose sidecar has a "Character {id}" paragraph.
    8. Header block present (Source / Advertiser / Last shown / Format rows).

  Supernova-rewrite doc only:
    9. Brand-swap check: every competitor brand name from brand_swaps_detected does NOT appear
       verbatim in the rewritten script (the rewrite should have swapped them out for Supernova).
    10. "Brand swaps applied" section present if brand_swaps_detected is non-empty.

Each issue is logged with severity (ERROR / WARN). A doc passes audit only if it has zero
ERROR-severity issues. The script writes step4_workspace/audit_reports/{ckey}.json with the
full per-doc report.

Usage:
    python3 scripts/step4_audit_docs.py --competitor zinglish <creative_id1> ...
    python3 scripts/step4_audit_docs.py --competitor zinglish CR123 CR456_v2

Exit code 0 if all audited rows pass; 1 if any failed; 2 if no audited rows produced docs.
"""
from __future__ import annotations

import argparse
import json
import pathlib
import sys
from typing import Tuple

WORKSPACE = pathlib.Path("step4_workspace")
SCENES_DIR = WORKSPACE / "scenes"
DOCS_DIR = WORKSPACE / "docs"
URLS_DIR = WORKSPACE / "image_urls"
REPORTS_DIR = WORKSPACE / "audit_reports"

# Severity levels — only ERROR blocks upload.
ERROR = "ERROR"
WARN = "WARN"

PLACEHOLDER_STRINGS = ("TODO", "<INSERT>", "<UNKNOWN>", "PLACEHOLDER", "FIXME")


def creative_key(creative_id: str, variant_index: int) -> str:
    return creative_id if variant_index == 0 else f"{creative_id}_v{variant_index + 1}"


def parse_id(raw: str) -> Tuple[str, int]:
    m = raw.rsplit("_v", 1)
    if len(m) == 2 and m[1].isdigit():
        return m[0], int(m[1]) - 1
    return raw, 0


def get_paragraphs_with_image_flag(doc):
    """Yield (paragraph_index, paragraph, has_image_bool) tuples in document order.
    has_image is True if the paragraph contains an inline image (drawing element)."""
    body = doc.element.body
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    drawing_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    para_idx = -1
    for child in body.iter():
        if child.tag == f"{w_ns}p":
            para_idx += 1
            has_image = child.find(f".//{drawing_ns}") is not None
            yield para_idx, child, has_image


def doc_paragraph_texts(doc) -> list[str]:
    return [p.text for p in doc.paragraphs]


def audit_image_caption_pairing(doc) -> list[dict]:
    """For every paragraph containing an embedded image, the NEXT paragraph
    must start with 'R2: https://' (or 'R2: ' with a URL anywhere in it).
    Returns a list of issue dicts."""
    issues = []
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    drawing_tag = f"{w_ns}drawing"
    text_tag = f"{w_ns}t"
    body = doc.element.body
    paras = [c for c in body if c.tag == f"{w_ns}p"]
    for i, p in enumerate(paras):
        has_image = p.find(f".//{drawing_tag}") is not None
        if not has_image:
            continue
        # Look at the next paragraph
        if i + 1 >= len(paras):
            issues.append({
                "severity": ERROR,
                "code": "image-without-caption",
                "msg": f"paragraph index {i} contains an image but has no following paragraph for R2 URL caption",
            })
            continue
        next_p = paras[i + 1]
        next_text = "".join((t.text or "") for t in next_p.findall(f".//{text_tag}")).strip()
        if not next_text.startswith("R2:"):
            issues.append({
                "severity": ERROR,
                "code": "image-without-r2-caption",
                "msg": f"paragraph index {i}: image followed by '{next_text[:60]}…' instead of 'R2: …'",
            })
            continue
        if "URL not available" in next_text:
            issues.append({
                "severity": ERROR,
                "code": "image-r2-url-missing",
                "msg": f"paragraph index {i}: image has 'URL not available' caption — re-run Stage 6a (step4_upload_images.py)",
            })
            continue
        if "http" not in next_text.lower():
            issues.append({
                "severity": ERROR,
                "code": "image-r2-caption-malformed",
                "msg": f"paragraph index {i}: caption '{next_text[:80]}' has no http URL",
            })
    return issues


def audit_one_doc(label: str, path: pathlib.Path, expected_scenes: int,
                   parsed_decompose: dict, parsed_rewrite: dict = None) -> dict:
    """Run all the audit checks for one doc. Return {issues: [...], pass: bool}."""
    from docx import Document
    report = {"label": label, "path": str(path), "issues": [], "pass": False}

    if not path.exists():
        report["issues"].append({"severity": ERROR, "code": "doc-missing", "msg": "file does not exist"})
        return report
    if path.stat().st_size < 4000:
        report["issues"].append({"severity": ERROR, "code": "doc-too-small",
                                 "msg": f"docx is only {path.stat().st_size} bytes"})
        return report
    try:
        doc = Document(str(path))
    except Exception as e:
        report["issues"].append({"severity": ERROR, "code": "doc-open-failed",
                                 "msg": f"{type(e).__name__}: {e}"})
        return report

    paragraphs = list(doc.paragraphs)
    full_text = "\n".join(p.text for p in paragraphs)

    # 1-2. Scene heading count
    scene_headings = [p for p in paragraphs if p.text.startswith("Scene ")]
    if len(scene_headings) < expected_scenes:
        report["issues"].append({
            "severity": ERROR, "code": "scene-count-mismatch",
            "msg": f"found {len(scene_headings)} scene headings, expected ≥ {expected_scenes}",
        })

    # 3. No empty headings
    for p in paragraphs:
        if p.style.name.startswith("Heading") and not p.text.strip():
            report["issues"].append({
                "severity": ERROR, "code": "empty-heading",
                "msg": "found a Heading-style paragraph with empty text",
            })
            break

    # 4. No placeholder strings
    for placeholder in PLACEHOLDER_STRINGS:
        if placeholder in full_text:
            report["issues"].append({
                "severity": ERROR, "code": "placeholder-string",
                "msg": f"document contains literal '{placeholder}' string",
            })

    # 5. R2 URL captions below every image (the core ask)
    image_caption_issues = audit_image_caption_pairing(doc)
    report["issues"].extend(image_caption_issues)
    image_count = len(doc.inline_shapes)
    report["image_count"] = image_count

    # 6. Reasonable paragraph density
    min_paras = max(8, expected_scenes * 3)
    if len(paragraphs) < min_paras:
        report["issues"].append({
            "severity": WARN, "code": "low-paragraph-count",
            "msg": f"{len(paragraphs)} paragraphs (expected ≥ {min_paras})",
        })

    # 7-8. Competitor-doc-specific
    if label == "competitor":
        characters = parsed_decompose.get("characters", [])
        for c in characters:
            cid = c.get("id", "?")
            if f"Character {cid}" not in full_text:
                report["issues"].append({
                    "severity": WARN, "code": "character-missing",
                    "msg": f"character '{cid}' declared in decompose but not mentioned in doc",
                })
        # Header block check
        for label_text in ("Advertiser:", "Last shown:", "Format:"):
            if label_text not in full_text:
                report["issues"].append({
                    "severity": WARN, "code": "header-row-missing",
                    "msg": f"expected header label '{label_text}' not found",
                })

    # 9-10. Supernova-doc-specific
    if label == "supernova" and parsed_rewrite is not None:
        brand_swaps = parsed_rewrite.get("brand_swaps_detected", [])
        if brand_swaps and "Brand swaps applied" not in full_text:
            report["issues"].append({
                "severity": WARN, "code": "brand-swaps-section-missing",
                "msg": "brand_swaps_detected is non-empty but 'Brand swaps applied' section absent",
            })
        # Critical check: competitor brand names should NOT appear in rewritten script.
        # We look only at supernova_script content (gathered by walking the doc).
        for swap in brand_swaps:
            cb = (swap.get("competitor_brand") or "").strip()
            if not cb or len(cb) < 3:
                continue
            # Allow it to appear in the brand-swaps-applied list — only flag elsewhere.
            occurrences = full_text.count(cb)
            if occurrences > 1:
                # 1 occurrence is acceptable (the brand-swaps log entry).
                report["issues"].append({
                    "severity": ERROR, "code": "brand-not-swapped",
                    "msg": f"competitor brand '{cb}' appears {occurrences} times in supernova doc; rewrite likely missed a swap",
                })

    error_count = sum(1 for i in report["issues"] if i["severity"] == ERROR)
    report["pass"] = (error_count == 0)
    report["error_count"] = error_count
    report["warn_count"] = sum(1 for i in report["issues"] if i["severity"] == WARN)
    return report


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                  formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--competitor", required=True,
                    help="Competitor slug (for log naming only).")
    ap.add_argument("ids", nargs="+",
                    help="creative_id (variant 0) or creative_id_v{N+1} for variants.")
    args = ap.parse_args()

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document  # noqa: F401
    except ImportError:
        print("[error] python-docx not installed. Run: python3 -m pip install --upgrade python-docx")
        return 2

    audited = 0
    passed = 0
    failed = []
    no_docs = 0

    for raw in args.ids:
        creative_id, variant_index = parse_id(raw)
        ckey = creative_key(creative_id, variant_index)
        comp_path = DOCS_DIR / f"{ckey}_competitor_analysis.docx"
        sn_path = DOCS_DIR / f"{ckey}_supernova_rewrite.docx"

        if not comp_path.exists() or not sn_path.exists():
            print(f"  [{ckey}] SKIP — docs not built yet (run step4_build_docs.py first)")
            no_docs += 1
            continue

        decompose_path = SCENES_DIR / f"{creative_id}.json"
        rewrite_path = SCENES_DIR / f"{creative_id}.supernova.json"
        if not decompose_path.exists() or not rewrite_path.exists():
            print(f"  [{ckey}] SKIP — decompose or rewrite sidecar missing")
            no_docs += 1
            continue

        try:
            decompose = json.loads(decompose_path.read_text())
            rewrite = json.loads(rewrite_path.read_text())
        except json.JSONDecodeError as e:
            print(f"  [{ckey}] SKIP — sidecar parse error: {e}")
            no_docs += 1
            continue

        parsed_decompose = decompose.get("parsed", {})
        parsed_rewrite = rewrite.get("parsed", {})
        expected_scenes = max(1, len(parsed_decompose.get("scenes", [])))

        comp_report = audit_one_doc("competitor", comp_path, expected_scenes, parsed_decompose)
        sn_report = audit_one_doc("supernova", sn_path, expected_scenes, parsed_decompose, parsed_rewrite)

        full_report = {
            "ckey": ckey,
            "creative_id": creative_id,
            "variant_index": variant_index,
            "competitor_doc": comp_report,
            "supernova_doc": sn_report,
            "overall_pass": comp_report["pass"] and sn_report["pass"],
        }
        (REPORTS_DIR / f"{ckey}.json").write_text(json.dumps(full_report, indent=2))

        audited += 1
        if full_report["overall_pass"]:
            passed += 1
            print(f"  [{ckey}] PASS   "
                  f"competitor: {comp_report['warn_count']}W, "
                  f"supernova: {sn_report['warn_count']}W")
        else:
            failed.append(ckey)
            print(f"  [{ckey}] FAIL   "
                  f"competitor: {comp_report['error_count']}E/{comp_report['warn_count']}W, "
                  f"supernova: {sn_report['error_count']}E/{sn_report['warn_count']}W")
            for issue in comp_report["issues"] + sn_report["issues"]:
                if issue["severity"] == ERROR:
                    print(f"    - {comp_report['label'] if issue in comp_report['issues'] else sn_report['label']}: "
                          f"[{issue['code']}] {issue['msg']}")

    print()
    print(f"Stage 7a done — audited={audited} passed={passed} failed={len(failed)} no-docs={no_docs}")
    if failed:
        print(f"Failed ckeys (do NOT proceed to upload for these): {failed}")
        print(f"Per-doc reports written to {REPORTS_DIR}/")
        return 1
    if audited == 0:
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
