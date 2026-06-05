#!/usr/bin/env python3
"""
Step 4 — Stage 6b (Google): build two .docx files per ad with R2 URL captions.

For each row, produces:
  step4_workspace/docs/<ckey>_competitor_analysis.docx
  step4_workspace/docs/<ckey>_supernova_rewrite.docx
where ckey = creative_id (variant 0) or {creative_id}_v{N+1} (variant N≥1).

Prereq: Stage 6a (step4_upload_images.py) must have run first. It writes
step4_workspace/image_urls/{ckey}.json which this script reads to look up the
R2 public URL for each embedded image. The URL is rendered as a small-italic
caption immediately below the picture so a reader can copy it directly from the
docx.

Format-aware:
  - YouTubeVideo: multi-scene docx with original frames + clean panels
  - Image:        single-page docx with the original image + clean panel
  - Text:         single-page docx with the headline + description verbatim (no images)

After build, programmatic verification runs (see verify_docs at the bottom of
this file). A separate Stage 7a (step4_audit_docs.py) does the stricter audit
including the R2-URL-below-every-image check; Stage 7b is the Claude review.

Usage:
    python3 scripts/step4_build_docs.py --competitor zinglish <creative_id1> ...
    python3 scripts/step4_build_docs.py --competitor zinglish CR123 CR456_v2
"""
from __future__ import annotations

import argparse
import csv
import json
import pathlib
import subprocess
import sys
from typing import Optional, Tuple

WORKSPACE = pathlib.Path("step4_workspace")
SCENES_DIR = WORKSPACE / "scenes"
FRAMES_DIR = WORKSPACE / "frames"
IMAGES_DIR = WORKSPACE / "images"
URLS_DIR = WORKSPACE / "image_urls"
DOCS_DIR = WORKSPACE / "docs"


# Caption text used when no R2 URL is available (image upload either failed or Stage 6a wasn't run).
NO_URL_CAPTION = "R2: (URL not available — re-run Stage 6a: step4_upload_images.py)"


def creative_key(creative_id: str, variant_index: int) -> str:
    """Matches the convention used by step4_upload_images.py."""
    return creative_id if variant_index == 0 else f"{creative_id}_v{variant_index + 1}"


def parse_id(raw: str) -> Tuple[str, int]:
    """Accept either '{creative_id}' or '{creative_id}_v{N+1}'."""
    m = raw.rsplit("_v", 1)
    if len(m) == 2 and m[1].isdigit():
        return m[0], int(m[1]) - 1
    return raw, 0


def load_image_url_lookup(ckey: str) -> dict:
    """Return {local_filename: r2_url} for all images of one creative.
    Reads step4_workspace/image_urls/{ckey}.json written by Stage 6a."""
    path = URLS_DIR / f"{ckey}.json"
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text())
    except (json.JSONDecodeError, OSError):
        return {}
    lookup = {}
    for arr_key in ("competitor_analysis_images", "supernova_rewrite_images"):
        for img in data.get(arr_key, []):
            fn = img.get("local_filename")
            url = img.get("r2_url")
            if fn and url:
                # Same filename may appear in both arrays (clean panels); same URL — keep one.
                lookup[fn] = url
    return lookup


def add_image_with_caption(doc, image_path: pathlib.Path, url_lookup: dict, log, *, width_inches: float = 4.5) -> bool:
    """Embed the image AND append an italic small-font 'R2: <url>' paragraph below.
    If no URL is found for the image, emit NO_URL_CAPTION so the gap is visible to the auditor.
    Returns True on success."""
    from docx.shared import Inches, Pt
    try:
        doc.add_picture(str(image_path), width=Inches(width_inches))
    except Exception as e:
        log(f"    image embed failed for {image_path}: {e}")
        return False
    url = url_lookup.get(image_path.name)
    cap_text = f"R2: {url}" if url else NO_URL_CAPTION
    cap_para = doc.add_paragraph(cap_text)
    for run in cap_para.runs:
        run.font.size = Pt(8)
        run.italic = True
    return True


def ensure_packages():
    try:
        import docx  # noqa
        import PIL  # noqa
    except ImportError:
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "--quiet",
             "python-docx", "Pillow"],
            check=False,
        )
    URLS_DIR.mkdir(parents=True, exist_ok=True)


def find_master_row(creative_id: str, competitor: str) -> Optional[dict]:
    master_path = pathlib.Path("master") / f"{competitor}.csv"
    if not master_path.exists():
        return None
    with master_path.open(encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            if (row.get("creative_id") or "").strip() == creative_id:
                return row
    return None


def video_path_for(creative_id: str, competitor: str) -> Optional[pathlib.Path]:
    candidates = sorted(pathlib.Path("videos").glob(f"{competitor}-*"))
    if not candidates:
        return None
    vdir = candidates[-1]
    p = vdir / f"{creative_id}.mp4"
    return p if p.exists() else None


def get_video_meta(vpath: pathlib.Path) -> dict:
    try:
        out = subprocess.check_output(
            ["ffprobe", "-v", "error", "-print_format", "json",
             "-show_format", "-show_streams", str(vpath)],
            timeout=10,
        ).decode()
        d = json.loads(out)
        dur = float(d.get("format", {}).get("duration", 0.0))
        w = h = None
        for s in d.get("streams", []):
            if s.get("codec_type") == "video" and w is None:
                w, h = s.get("width"), s.get("height")
        return {"duration_s": round(dur, 1),
                "resolution": f"{w}x{h}" if w else "?"}
    except Exception:
        return {"duration_s": None, "resolution": "?"}


def panel_image_path(creative_id: str, scene_n: int, position: str) -> Optional[pathlib.Path]:
    p = IMAGES_DIR / creative_id / f"gen_{scene_n:02d}_{position}.jpg"
    return p if p.exists() and p.stat().st_size > 30000 else None


def original_frame_path(creative_id: str, scene_n: int) -> Optional[pathlib.Path]:
    p = FRAMES_DIR / creative_id / f"orig_{scene_n:02d}.jpg"
    return p if p.exists() and p.stat().st_size > 5000 else None


def add_header_block(doc, creative_id: str, competitor: str, fmt: str,
                      master_row: Optional[dict], parsed: dict, decompose: dict) -> None:
    from docx.shared import Pt
    advertiser = (master_row or {}).get("advertiser_name", "?")
    regions = (master_row or {}).get("regions_shown", "?")
    # Google's Transparency Center only exposes "Last shown" for India ads; no "First shown".
    last_shown = (master_row or {}).get("creative_last_shown_date", "?")

    doc.add_paragraph(f"Source: Google Ads Transparency Center")
    doc.add_paragraph(f"Advertiser: {advertiser}   |   Regions: {regions}")
    doc.add_paragraph(f"Last shown: {last_shown}")
    doc.add_paragraph(f"Format: {fmt}")
    pt = parsed.get("production_type", "?")
    doc.add_paragraph(f"Production type: {pt}")
    if parsed.get("production_notes"):
        doc.add_paragraph(f"Production notes: {parsed['production_notes']}")


def build_competitor_doc(creative_id: str, variant_index: int, competitor: str,
                          decompose: dict, out_path: pathlib.Path, log) -> bool:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    parsed = decompose.get("parsed", {})
    fmt = decompose.get("creative_format", "YouTubeVideo")
    scenes = parsed.get("scenes", [])
    characters = parsed.get("characters", [])
    master_row = find_master_row(creative_id, competitor)

    # Load R2 URL lookup written by Stage 6a (step4_upload_images.py).
    ckey = creative_key(creative_id, variant_index)
    url_lookup = load_image_url_lookup(ckey)

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)

    variant_suffix = f" (variant {variant_index + 1}/{variant_index + 1}+)" if variant_index > 0 else ""
    title = doc.add_heading(
        f"Competitor Ad Analysis — {competitor.title()} — {fmt} — {creative_id}{variant_suffix}",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_header_block(doc, creative_id, competitor, fmt, master_row, parsed, decompose)

    if fmt == "YouTubeVideo":
        vpath = video_path_for(creative_id, competitor)
        meta = get_video_meta(vpath) if vpath else {"duration_s": None, "resolution": "?"}
        doc.add_paragraph(f"Source video: {vpath.name if vpath else '?'}")
        doc.add_paragraph(f"Duration: {meta.get('duration_s', '?')}s   |   "
                          f"Resolution: {meta.get('resolution', '?')}")

    # Character roster + reference sheet images (video only)
    if fmt == "YouTubeVideo":
        doc.add_heading("Character roster", level=2)
        if not characters:
            doc.add_paragraph("(no characters detected)")
        for c in characters:
            cid = c.get("id", "?")
            role = c.get("role", "")
            doc.add_paragraph(f"Character {cid} — {role}", style="List Bullet")
            if c.get("appearance"):
                doc.add_paragraph(f"    Appearance: {c['appearance']}")
            if c.get("wardrobe"):
                doc.add_paragraph(f"    Wardrobe: {c['wardrobe']}")
            if c.get("demeanor"):
                doc.add_paragraph(f"    Demeanor: {c['demeanor']}")
            # Embed the character reference sheet if it exists.
            # Step 3 (character_sheets) writes to images/{creative_id}/, not images/{ckey}/.
            char_sheet = IMAGES_DIR / creative_id / f"char_{cid}_sheet.jpg"
            if char_sheet.exists() and char_sheet.stat().st_size > 50_000:
                doc.add_paragraph(f"    Reference sheet:")
                add_image_with_caption(doc, char_sheet, url_lookup, log, width_inches=3.5)
        doc.add_paragraph()
        doc.add_paragraph("─" * 60)

    # Per-scene blocks (works for video AND single-scene image/text)
    for scene in scenes:
        n = scene.get("n", 0)
        label = scene.get("scene_label", "?")
        doc.add_heading(f"Scene {n} — {label}", level=1)

        if fmt == "YouTubeVideo":
            ss = scene.get("start_s", 0)
            es = scene.get("end_s", 0)
            doc.add_paragraph(f"[{ss}s – {es}s]   Camera: {scene.get('camera', '?')}   "
                              f"Layout: {scene.get('layout', '?')}")
        if scene.get("setting"):
            doc.add_paragraph(f"Setting: {scene['setting']}")

        # Original image / frame
        orig = original_frame_path(creative_id, n)
        if orig:
            doc.add_paragraph("Original:")
            add_image_with_caption(doc, orig, url_lookup, log)
        elif fmt != "Text":
            doc.add_paragraph("(no original screenshot extracted)")

        # Clean panels (video + image)
        panels = scene.get("panels", []) if fmt != "Text" else []
        if panels:
            doc.add_paragraph("Clean regenerated panels:")
            for panel in panels:
                pos = panel.get("position", "full")
                pp = panel_image_path(creative_id, n, pos)
                if pp:
                    doc.add_paragraph(f"  Panel: {pos}")
                    add_image_with_caption(doc, pp, url_lookup, log)
                else:
                    doc.add_paragraph(f"  Panel: {pos} — (regeneration unavailable; "
                                       f"content-policy block or pending)")

        # Visual description
        if scene.get("visual_description"):
            doc.add_heading("Composite visual description", level=3)
            doc.add_paragraph(scene["visual_description"])

        # Per-panel descriptions
        if panels:
            doc.add_heading("Per-panel descriptions", level=3)
            for panel in panels:
                pos = panel.get("position", "full")
                desc = panel.get("panel_visual_description", "")
                doc.add_paragraph(f"[{pos}] {desc}", style="List Bullet")

        # Audio transcript (video) OR description (text)
        if scene.get("audio_transcript"):
            heading = "Audio transcript" if fmt == "YouTubeVideo" else "Ad description"
            doc.add_heading(heading, level=3)
            for line in scene["audio_transcript"].split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # On-screen text / headline
        if scene.get("on_screen_text"):
            heading = "On-screen text" if fmt != "Text" else "Headline"
            doc.add_heading(heading, level=3)
            doc.add_paragraph(scene["on_screen_text"])

        # Text-only: messaging analysis
        if fmt == "Text" and parsed.get("messaging_analysis"):
            doc.add_heading("Messaging analysis", level=3)
            doc.add_paragraph(parsed["messaging_analysis"])

        doc.add_paragraph("─" * 60)

    # Provenance
    doc.add_heading("Provenance", level=2)
    doc.add_paragraph(f"Decomposed with: {decompose.get('model', '?')}")
    if fmt != "Text":
        doc.add_paragraph(f"Image generation: gemini-3-pro-image-preview (Nano Banana Pro)")
    doc.add_paragraph(f"This document is for internal competitor research. "
                       f"Sourced from Google Ads Transparency Center.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return True


def build_supernova_doc(creative_id: str, variant_index: int, competitor: str,
                         decompose: dict, rewrite: dict, out_path: pathlib.Path, log) -> bool:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    parsed_decompose = decompose.get("parsed", {})
    fmt = decompose.get("creative_format", "YouTubeVideo")
    parsed_rewrite = rewrite.get("parsed", {})
    scenes_d = {s.get("n"): s for s in parsed_decompose.get("scenes", [])}
    scenes_r = parsed_rewrite.get("scenes", [])
    brand_swaps = parsed_rewrite.get("brand_swaps_detected", [])

    # Load R2 URL lookup written by Stage 6a.
    ckey = creative_key(creative_id, variant_index)
    url_lookup = load_image_url_lookup(ckey)

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)

    variant_suffix = f" (variant {variant_index + 1})" if variant_index > 0 else ""
    title = doc.add_heading(
        f"Supernova-Voice Rewrite — based on {competitor.title()} {fmt} ad {creative_id}{variant_suffix}",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("This is an internal brand-safe rewrite of a competitor's ad, "
                       "produced for Supernova's creative team. Use the visuals + script "
                       "as a starting point for original Supernova creative — do not ship "
                       "this as-is.")
    doc.add_paragraph()
    doc.add_paragraph("─" * 60)

    # Per-scene blocks
    for scene_r in scenes_r:
        n = scene_r.get("n", 0)
        scene_d = scenes_d.get(n, {})
        label = scene_r.get("scene_label") or scene_d.get("scene_label", "?")

        doc.add_heading(f"Scene {n} — {label}", level=1)

        # Clean panel imagery (video + image; not text)
        if fmt != "Text":
            panels = scene_d.get("panels", [])
            for panel in panels:
                pos = panel.get("position", "full")
                pp = panel_image_path(creative_id, n, pos)
                if pp:
                    add_image_with_caption(doc, pp, url_lookup, log)

        # Scene summary
        if scene_r.get("scene_summary"):
            doc.add_heading("Scene summary", level=3)
            for line in scene_r["scene_summary"].split("\n"):
                line = line.strip().lstrip("- *•").strip()
                if line:
                    doc.add_paragraph(line, style="List Bullet")

        # Supernova rewrite script
        if scene_r.get("supernova_script"):
            heading = "Supernova-voice script" if fmt == "YouTubeVideo" else "Supernova-voice rewrite"
            doc.add_heading(heading, level=3)
            for line in scene_r["supernova_script"].split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # On-screen text rewrite
        if scene_r.get("supernova_on_screen_text"):
            heading = "On-screen text (Supernova version)" if fmt != "Text" else "Headline (Supernova version)"
            doc.add_heading(heading, level=3)
            doc.add_paragraph(scene_r["supernova_on_screen_text"])

        doc.add_paragraph("─" * 60)

    # Brand swaps log
    if brand_swaps:
        doc.add_heading("Brand swaps applied", level=2)
        for swap in brand_swaps:
            cb = swap.get("competitor_brand", "?")
            sw = swap.get("supernova_swap", "?")
            doc.add_paragraph(f"{cb} → {sw}", style="List Bullet")

    doc.add_heading("Provenance", level=2)
    doc.add_paragraph(f"Rewritten with: {rewrite.get('model', '?')}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return True


def verify_docs(creative_id: str, comp_path: pathlib.Path, sn_path: pathlib.Path,
                 expected_scenes: int, fmt: str) -> list[str]:
    """Basic structural verification. The fuller R2-URL-below-every-image audit
    runs in Stage 7a (step4_audit_docs.py)."""
    from docx import Document
    issues = []

    # Image/Text rows have 1 scene → looser thresholds
    min_paras_competitor = max(8, expected_scenes * 3)
    min_paras_supernova = max(5, expected_scenes * 2)

    for label, path, min_paras in (("competitor", comp_path, min_paras_competitor),
                                     ("supernova", sn_path, min_paras_supernova)):
        if not path.exists():
            issues.append(f"{label}: file missing")
            continue
        if path.stat().st_size < 4000:
            issues.append(f"{label}: too small ({path.stat().st_size}B)")
            continue
        try:
            doc = Document(str(path))
        except Exception as e:
            issues.append(f"{label}: failed to open: {e}")
            continue
        paragraphs = list(doc.paragraphs)
        if len(paragraphs) < min_paras:
            issues.append(f"{label}: only {len(paragraphs)} paragraphs (expected ≥ {min_paras})")
        scene_headings = [p for p in paragraphs if p.text.startswith("Scene ")]
        if len(scene_headings) < expected_scenes:
            issues.append(f"{label}: only {len(scene_headings)} scene headings (expected ≥ {expected_scenes})")
        for p in paragraphs:
            if p.style.name.startswith("Heading") and not p.text.strip():
                issues.append(f"{label}: empty heading found")
                break
        full_text = "\n".join(p.text for p in paragraphs)
        for placeholder in ("TODO", "<INSERT>", "<UNKNOWN>", "PLACEHOLDER"):
            if placeholder in full_text:
                issues.append(f"{label}: contains placeholder string '{placeholder}'")
        # Heuristic: count images vs. R2: captions. They should match.
        image_count = len(doc.inline_shapes)
        r2_caption_count = sum(1 for p in paragraphs if p.text.strip().startswith("R2:"))
        if image_count > 0 and r2_caption_count == 0:
            issues.append(f"{label}: {image_count} embedded images but zero 'R2:' captions — Stage 6a likely not run")
        if image_count > 0 and r2_caption_count < image_count:
            issues.append(f"{label}: {image_count} images but only {r2_caption_count} 'R2:' captions")

    return issues


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                  formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--competitor", required=True)
    ap.add_argument("ids", nargs="+",
                    help="creative_id (variant 0) or creative_id_v{N+1} for variants.")
    args = ap.parse_args()

    ensure_packages()
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    competitor = args.competitor.lower().strip()

    def log(msg):
        print(msg, flush=True)

    print(f"Stage 6b — building docx files for {len(args.ids)} ads")
    print(f"  (Prereq: Stage 6a step4_upload_images.py should have been run first.)")
    built = []
    failed_verify = []

    for raw in args.ids:
        creative_id, variant_index = parse_id(raw)
        ckey = creative_key(creative_id, variant_index)
        sidecar = SCENES_DIR / f"{creative_id}.json"
        rewrite_sc = SCENES_DIR / f"{creative_id}.supernova.json"
        if not sidecar.exists():
            log(f"  [{ckey}] SKIP — no decompose sidecar")
            continue
        if not rewrite_sc.exists():
            log(f"  [{ckey}] SKIP — no Supernova rewrite sidecar")
            continue
        decompose = json.loads(sidecar.read_text())
        rewrite = json.loads(rewrite_sc.read_text())
        fmt = decompose.get("creative_format", "YouTubeVideo")
        n_scenes = len(decompose.get("parsed", {}).get("scenes", []))

        # Doc filenames are keyed on ckey so variants get distinct files
        comp_path = DOCS_DIR / f"{ckey}_competitor_analysis.docx"
        sn_path = DOCS_DIR / f"{ckey}_supernova_rewrite.docx"

        # Warn if the URL sidecar is missing (means Stage 6a hasn't run)
        if not (URLS_DIR / f"{ckey}.json").exists():
            log(f"  [{ckey}] WARN — no image_urls/{ckey}.json sidecar; doc captions will show 'URL not available'. Run Stage 6a (step4_upload_images.py) first.")

        try:
            build_competitor_doc(creative_id, variant_index, competitor, decompose, comp_path, log)
            build_supernova_doc(creative_id, variant_index, competitor, decompose, rewrite, sn_path, log)
            issues = verify_docs(ckey, comp_path, sn_path, n_scenes, fmt)
            if issues:
                log(f"  [{ckey}] BUILT with verification issues:")
                for iss in issues:
                    log(f"    - {iss}")
                failed_verify.append((ckey, issues))
            else:
                log(f"  [{ckey}] OK — basic verification passed (format={fmt}, scenes={n_scenes})")
                built.append(ckey)
        except Exception as e:
            log(f"  [{ckey}] BUILD FAILED: {type(e).__name__}: {e}")

    print(f"\nStage 6b done — {len(built)} clean, {len(failed_verify)} with verification issues")
    print(f"NEXT: run step4_audit_docs.py for the strict R2-URL-below-every-image audit.")
    return 0 if not failed_verify else 1


if __name__ == "__main__":
    raise SystemExit(main())
